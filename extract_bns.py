"""
Extract BNS sections from official Word document and create structured JSON
"""
import json
import re
from docx import Document

def extract_bns_sections(docx_path):
    """Extract BNS sections from Word document"""
    doc = Document(docx_path)

    sections = []
    current_section = None
    current_text = []
    in_section = False

    print(f"Processing {len(doc.paragraphs)} paragraphs...")

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if not text:
            continue

        # Look for section numbers (e.g., "1.", "2.", "100." at start of line)
        # BNS sections are numbered 1-358
        section_match = re.match(r'^(\d{1,3})\.(\s+.*)?$', text)

        if section_match and int(section_match.group(1)) <= 500:
            # Save previous section
            if current_section:
                current_section['full_text'] = ' '.join(current_text).strip()
                sections.append(current_section)

            # Start new section
            section_num = section_match.group(1)
            section_title = section_match.group(2).strip() if section_match.group(2) else ""

            current_section = {
                'section_number': section_num,
                'title': section_title,
                'description': '',
                'punishment': '',
                'category': 'other',
                'severity': 'moderate',
                'keywords': [],
                'essential_elements': {}
            }
            current_text = [section_title] if section_title else []
            in_section = True

            print(f"Found section {section_num}: {section_title[:60]}...")

        elif in_section and current_section:
            # Accumulate section text
            current_text.append(text)

    # Save last section
    if current_section:
        current_section['full_text'] = ' '.join(current_text).strip()
        sections.append(current_section)

    print(f"\n✅ Extracted {len(sections)} sections")

    # Post-process sections to extract descriptions and categorize
    processed_sections = []

    for section in sections:
        # Extract description (usually the main definition/offense description)
        full_text = section['full_text']

        # Try to find punishment info
        punishment_patterns = [
            r'shall be punished with (.+?)(?:\.|$)',
            r'imprisonment.*?for.*?(?:term|period).*?(?:of|which may extend to)\s+([^.]+)',
            r'fine.*?(?:which may extend to|of)\s+([^.]+)'
        ]

        punishment = "To be determined by court"
        for pattern in punishment_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                punishment = match.group(0).strip()
                break

        # Basic categorization based on keywords
        category = categorize_section(full_text)
        severity = determine_severity(full_text)
        keywords = extract_keywords(full_text)

        processed_section = {
            'section_number': section['section_number'],
            'title': section['title'][:200],  # Limit title length
            'description': full_text[:1000],  # First 1000 chars as description
            'essential_elements': extract_elements(full_text),
            'punishment': punishment[:500],
            'category': category,
            'severity': severity,
            'keywords': keywords[:20]  # Top 20 keywords
        }

        processed_sections.append(processed_section)

    return processed_sections


def categorize_section(text):
    """Categorize crime based on keywords"""
    text_lower = text.lower()

    # Category keywords
    categories = {
        'violence': ['murder', 'assault', 'hurt', 'grievous', 'death', 'kill', 'attack', 'violence', 'injury'],
        'sexual': ['rape', 'sexual', 'modesty', 'outraging', 'intercourse', 'unnatural'],
        'property': ['theft', 'robbery', 'burglary', 'dacoity', 'extortion', 'criminal breach of trust', 'stolen'],
        'fraud': ['cheating', 'forgery', 'counterfeiting', 'fraudulent', 'dishonestly'],
        'cyber': ['cyber', 'computer', 'electronic', 'internet', 'digital'],
        'public_order': ['unlawful assembly', 'riot', 'affray', 'public nuisance', 'sedition'],
        'economic': ['criminal breach of trust', 'misappropriation', 'embezzlement']
    }

    scores = {}
    for category, keywords in categories.items():
        score = sum(1 for keyword in keywords if keyword in text_lower)
        if score > 0:
            scores[category] = score

    if scores:
        return max(scores, key=scores.get)
    return 'other'


def determine_severity(text):
    """Determine severity based on punishment mentioned"""
    text_lower = text.lower()

    if any(word in text_lower for word in ['death', 'life imprisonment', 'imprisonment for life']):
        return 'very_serious'
    elif any(word in text_lower for word in ['ten years', 'twenty years', 'fourteen years']):
        return 'serious'
    elif any(word in text_lower for word in ['imprisonment', 'rigorous']):
        return 'moderate'
    else:
        return 'minor'


def extract_keywords(text):
    """Extract important keywords from section text"""
    # Remove common words
    stop_words = {'the', 'is', 'at', 'which', 'on', 'a', 'an', 'and', 'or', 'but', 'in', 'with', 'to', 'for', 'of', 'as', 'by', 'from', 'be', 'shall', 'may', 'any', 'such', 'if', 'whoever', 'being', 'having'}

    # Extract words
    words = re.findall(r'\b[a-z]{4,}\b', text.lower())

    # Filter and count
    word_freq = {}
    for word in words:
        if word not in stop_words:
            word_freq[word] = word_freq.get(word, 0) + 1

    # Return top keywords
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    return [word for word, _ in sorted_words[:20]]


def extract_elements(text):
    """Extract essential legal elements from section"""
    elements = {}

    text_lower = text.lower()

    # Check for mens rea (mental state)
    if 'intentionally' in text_lower or 'knowingly' in text_lower:
        elements['mens_rea'] = 'intentional'
    elif 'rash' in text_lower or 'negligent' in text_lower:
        elements['mens_rea'] = 'rash/negligent'
    elif 'knowledge' in text_lower:
        elements['mens_rea'] = 'knowledge'

    # Check for common elements
    if 'dishonestly' in text_lower:
        elements['dishonest_intention'] = True

    if 'fraudulent' in text_lower:
        elements['fraudulent_intention'] = True

    if 'voluntarily' in text_lower:
        elements['voluntary_act'] = True

    return elements


if __name__ == "__main__":
    sections = extract_bns_sections("Full BNS.docx")

    # Save to JSON
    output_file = "bns_sections.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(sections, f, ensure_ascii=False, indent=2)

    print(f"\n✅ Saved {len(sections)} sections to {output_file}")

    # Show sample
    if sections:
        print("\n📄 Sample section:")
        print(json.dumps(sections[0], indent=2, ensure_ascii=False))
